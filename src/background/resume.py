#(c) 2024 Daniel DeMoney. All rights reserved.
from datetime import datetime
from typing import Dict
from mysql.connector.types import RowItemType
import zlib
from tika import parser
import docx
import io
import json
import logging

class Resume:
    '''
    Resume

    Object to represent the resumes users upload.

    id: auto-incremented id of the resume in the db
    user_id: uuid of user the resume is connected to
    file_name: name of the uploaded resume file
    file_type: type of file uploaded (extension, .docx for word)
    file_content: bytes, content of the resume (not a readable str)
    file_text: str, readable text of the resume
    upload_date: datetime of the upload
    '''
    def __init__(self, id: int, user_id: str, file_name: str, file_type: str, file_content: bytes, upload_date: datetime, file_text=None, name=None, isDefault=False) -> None:
        self.id = id
        self.name = name
        self.user_id = user_id
        self.file_name = file_name
        self.file_type = file_type
        self.file_content = file_content
        self.isDefault = isDefault
        if file_text:
            self.file_text = file_text
        elif self.file_type == "pdf":
            #front end will never know the text until we process the pd
            result = parser.from_buffer(file_content)
            self.file_text = result["content"]
        elif self.file_type == "docx":
            #Asking for forgiveness
            try:
                mock_file_descriptor = io.BytesIO(file_content)
                doc = docx.Document(mock_file_descriptor)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                self.file_text = "\n".join(full_text)
            except TypeError:
                doc = docx.Document(file_content)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                self.file_text = "\n".join(full_text)
        self.upload_date = upload_date
    '''
    compress

    quick alogirithm for str to bytes compression, used for the text of the resume

    We just use the zlib base compress decompress algos for the resume content, (raw bytes of file)

    resume_str: text of resume

    returns:

    bytes: compressed text of resume
    '''
    def compress(resume_str: str) -> bytes:
        resume_bytes: bytes = resume_str.encode("utf-8")
        compressed_resume: bytes = zlib.compress(resume_bytes)
        return compressed_resume
    '''
    decompress

    reverse alogirthm for the above function to decompress our compressed resume text.

    compressed_resume: the compressed version of the resumes text

    retutns:

    uncompressed resume
    '''
    def decompress(compressed_resume: bytes) -> str:
        decompressed_bytes: bytes = zlib.decompress(compressed_resume)
        decompressed_resume: str = decompressed_bytes.decode("utf-8")
        return decompressed_resume
    '''
    create_with_sql_row

    creates a resume from a sql row

    sql_query_row: result of cursor.fectchone

    returns:

    resume
    '''
    @classmethod
    def create_with_sql_row(cls, sql_query_row: (Dict[str, RowItemType])) -> 'Resume':
        logging.info("CREATING RESUME WITH SQL ROW")
        id: int = sql_query_row["Id"]
        user_id: str = sql_query_row["UserId"]
        name: str | None = sql_query_row["Name"]
        file_name: str = sql_query_row["FileName"]
        file_type: str = sql_query_row["FileType"]
        file_content: bytes = zlib.decompress(sql_query_row["FileContent"])
        file_text: str = Resume.decompress(sql_query_row["FileText"])
        upload_date: datetime = sql_query_row["UploadDate"]
        isDefault: bool = sql_query_row["IsDefault"]
        return cls(id, user_id, file_name, file_type, file_content, upload_date, file_text=file_text, name=name, isDefault=isDefault)
    '''
    create_with_json

    creates a resume based off the json object sent from client

    json:

    json object sent from client

    returns:

    resume
    '''
    @classmethod
    def create_with_json(cls, jsonObj: Dict) -> 'Resume':
        logging.info("CREATING RESUME WITH JSON")
        filtered_data = {k: v for k, v in jsonObj.items() if k not in ['fileContent', 'fileText']}
        logging.info(json.dumps(filtered_data, indent=2))
        try:
            id: int = jsonObj["id"]
        except KeyError:
            id: int = None
        user_id: str = jsonObj["userId"]
        name: str | None = jsonObj["name"]
        file_name: str = jsonObj["fileName"]
        file_type: str = jsonObj["fileType"]
        file_content: bytes = bytes(jsonObj["fileContent"])
        isDefault: bool = jsonObj["isDefault"]
        return cls(id, user_id, file_name, file_type, file_content, None, name=name, isDefault=isDefault)
    '''
    to_json

    dumps a resume to json to be sent back to the client

    returns:

    json
    '''
    def to_json(self) -> Dict:
        return {
            "id": self.id,
            "userId": self.user_id,
            "name": self.name,
            "fileName": self.file_name,
            "fileType": self.file_type,
            "fileContent": list(self.file_content),
            "fileText": self.file_text,
            "uploadDate": int(self.upload_date.timestamp()) if self.upload_date else None,
            "isDefault": self.isDefault
        }
    '''
    to_sql_friendly_json

    dumps a resume to json to be added to db, compresses necessary data

    returns:

    json
    '''
    def to_sql_friendly_json(self) -> Dict:
        return {
            "id": self.id,
            "userId": self.user_id,
            "name": self.name,
            "fileName": self.file_name,
            "fileType": self.file_type,
            "fileContent": zlib.compress(self.file_content),
            "fileText": Resume.compress(self.file_text),
            "isDefault": self.isDefault
        }
    